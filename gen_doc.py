"""
Generate comprehensive project documentation as a DOCX file.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml.ns as nsmap

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Styles helper ─────────────────────────────────────────────────────────────
def set_style(para, bold=False, size=12, color=None, italic=False):
    for run in para.runs:
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    return p

def h4(text):
    p = doc.add_heading(text, level=4)
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'F3F4F6')
    p._p.get_or_add_pPr().append(shading)
    return p

def note(text):
    p = doc.add_paragraph()
    run = p.add_run("ℹ  " + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x37, 0x51, 0x8c)
    return p

def divider():
    doc.add_paragraph('─' * 80)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("POWER PROFILE SYSTEM")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Graduation Project — Complete Technical Documentation")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x4b, 0x5e, 0x80)

doc.add_paragraph()
doc.add_paragraph()

meta_lines = [
    ("System Type",    "Building Energy Management & Power Analytics Platform"),
    ("Backend",        "Laravel 11 (PHP 8.2) — RESTful API with Sanctum Auth"),
    ("Frontend",       "React 18 + Tailwind CSS — Single Page Application"),
    ("Database",       "MySQL — Relational with Polymorphic Relationships"),
    ("External APIs",  "NASA POWER (Solar Irradiance), Google OAuth 2.0"),
    ("Standards",      "IEC 60364-8-1, NEC/IEC Motor Inrush, PENRA Power Factor"),
]
t = doc.add_table(rows=len(meta_lines), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(meta_lines):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    "1.  Project Overview & Purpose",
    "2.  System Architecture",
    "3.  Database Schema & Entity Relationships",
    "4.  Backend — Laravel Application",
    "    4.1  Authentication & Authorization",
    "    4.2  Project Hierarchy (Project → Building → Floor → Room)",
    "    4.3  Electrical Components System",
    "    4.4  Power Sources — Utility, Generator, Sockets",
    "    4.5  Total Power Controller — Diversity Factors & Vector Power",
    "    4.6  Load Profile Controller — Hourly Demand Curves",
    "    4.7  Schedule Controller — Source Dispatch",
    "    4.8  Phase Balance Controller — 3-Phase Optimization",
    "    4.9  Solar Irradiance Service — NASA POWER Integration",
    "    4.10 Source Dispatch Service",
    "    4.11 Socket Demand Service",
    "    4.12 Backup & Restore System",
    "    4.13 Project Member Management",
    "5.  Frontend — React Application",
    "    5.1  Routing & Authentication",
    "    5.2  Dashboard Page",
    "    5.3  Project Page",
    "    5.4  Building / Floor / Room Pages",
    "    5.5  Load Schedule Page",
    "    5.6  Phase Balance Page",
    "    5.7  Power Banner Component",
    "    5.8  Power Sources Banner Component",
    "6.  API Reference — All Endpoints",
    "7.  Electrical Engineering Theory",
    "    7.1  Apparent / Active / Reactive Power & Power Factor",
    "    7.2  IEC 60364-8-1 Diversity Factors",
    "    7.3  Motor Inrush Current (NEC/IEC 125% Rule)",
    "    7.4  Reactive Power Correction & Capacitor Bank Sizing",
    "    7.5  Phase Imbalance & Neutral Current",
    "    7.6  Solar PV — PSH, STC, Performance Ratio",
    "    7.7  Priority-Based Load Classification",
    "8.  Calculation Walkthroughs (End-to-End Examples)",
    "9.  Known Limitations & Future Work",
]
for item in toc_items:
    body(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Project Overview & Purpose")

body(
    "Power Profile is a full-stack web application built as a graduation project. It solves the "
    "problem of designing and analysing the electrical power demand of multi-building complexes — "
    "universities, hospitals, industrial parks, or any facility that contains many buildings, floors, "
    "and rooms each with their own electrical loads."
)
body(
    "A structural engineer or electrical designer can create a digital twin of their facility "
    "inside the system: they define buildings, floors, rooms, and then add electrical components "
    "(air conditioners, servers, lighting, motors, etc.) at any level of the hierarchy. The system "
    "then automatically calculates:"
)
bullet("The total apparent power demand (VA) and active power (W) the facility will draw at peak.")
bullet("A 24-hour hourly load profile showing when each kW of demand occurs throughout the day.")
bullet("How much solar energy a rooftop PV system can contribute, fetched live from NASA satellite data.")
bullet("How solar, utility grid, and diesel generator sources should be dispatched hour-by-hour.")
bullet("The reactive power (kVAR) present and how to correct it with a capacitor bank.")
bullet("Whether the three-phase electrical supply is balanced across phases A, B, and C.")
bullet("The critical, essential, and normal load breakdown for UPS and emergency sizing.")
body(
    "The application is designed to cover the full electrical feasibility workflow from initial "
    "load estimation to source sizing and power quality analysis — all in one tool."
)

doc.add_paragraph()
h2("Core Design Goals")
bullet("Hierarchical data model (Project → Building → Floor → Room → Component) mirrors real construction practice.")
bullet("All power calculations follow published electrical standards: IEC 60364-8-1, PENRA, NEC.")
bullet("Solar data is real satellite-derived irradiance from NASA POWER, not estimated averages.")
bullet("Critical loads are never diversified — they count at full demand always.")
bullet("The application is multi-user with role-based access (admin / main / normal).")
bullet("Full backup and restore at every hierarchy level.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. System Architecture")

h2("2.1 Technology Stack")
t = doc.add_table(rows=8, cols=2)
t.style = 'Table Grid'
rows_data = [
    ("Layer",          "Technology"),
    ("Backend Framework", "Laravel 11 — PHP 8.2"),
    ("API Style",      "RESTful JSON API, all responses as application/json"),
    ("Authentication", "Laravel Sanctum (token-based) + Google OAuth 2.0 via Socialite"),
    ("Frontend",       "React 18, React Router v6, Tailwind CSS 3"),
    ("HTTP Client",    "Axios (frontend) / Guzzle via Laravel HTTP facade (backend → NASA)"),
    ("Database",       "MySQL 8 — InnoDB, with JSON column type for arrays"),
    ("Caching",        "Laravel Cache (file/redis driver) — NASA API results cached 30 days"),
]
for i, (k, v) in enumerate(rows_data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
h2("2.2 Overall Architecture")
body(
    "The backend is a stateless RESTful API. Every request carries a Bearer token issued by "
    "Sanctum. The frontend is a React SPA that fetches data exclusively through the API. "
    "There is no server-side rendering."
)
code_block(
    "Browser (React SPA)\n"
    "      │  HTTPS JSON API\n"
    "      ▼\n"
    "Laravel 11 API Server\n"
    "      │\n"
    "      ├── MySQL database  (projects, components, sockets, …)\n"
    "      ├── Laravel Cache   (NASA POWER results, 30-day TTL)\n"
    "      └── NASA POWER API  (satellite solar irradiance, external HTTP)"
)

h2("2.3 Request Lifecycle")
body("Every authenticated API request follows this path:")
bullet("1. Sanctum middleware validates the Bearer token from the Authorization header.")
bullet("2. throttle:api-general (60 req/min) or throttle:api-heavy (20 req/min) rate limiting is applied.")
bullet("3. Route model binding automatically resolves {project}, {building}, {floor}, {room} IDs into Eloquent model instances.")
bullet("4. The controller method runs a userRole() check — if the authenticated user has no role on the project, a 403 Forbidden is returned immediately.")
bullet("5. Business logic executes and a JSON response is returned.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Database Schema & Entity Relationships")

h2("3.1 Hierarchy Tables")
body(
    "The core data model forms a strict parent-child tree. Every level can have its own "
    "components, utility lines, generator lines, and sockets."
)

h3("projects")
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Column"; t.rows[0].cells[1].text = "Type"; t.rows[0].cells[2].text = "Description"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
proj_cols = [
    ("id",                     "bigint PK",  "Auto-increment primary key"),
    ("user_id",                "FK → users", "Owner user (always has 'admin' role)"),
    ("name",                   "varchar",    "Project display name"),
    ("building_type",          "varchar",    "Type of buildings (university, hospital, etc.)"),
    ("buildings_count",        "int",        "Expected number of buildings (informational)"),
    ("total_power",            "decimal",    "Cached total power (legacy field)"),
    ("solar_power",            "decimal",    "Installed PV capacity in kW"),
    ("existing_solar_power",   "decimal",    "User-declared existing PV in W (when solar_source='existing')"),
    ("solar_source",           "varchar",    "'max' (auto-compute from roof area) or 'existing' (user-declared)"),
    ("generator_source",       "varchar",    "'max' or 'existing'"),
    ("generator_power",        "decimal",    "Generator capacity in VA"),
    ("auto_backup_interval",   "int",        "Minutes between auto-backups (0 = off)"),
    ("last_auto_backup_at",    "datetime",   "Timestamp of last auto-backup"),
    ("work_days",              "JSON array", "e.g. ['monday','tuesday',...] — entity-level schedule"),
    ("work_time_intervals",    "JSON array", "Daily work time windows"),
    ("working_season_intervals","JSON array","Active months, e.g. [{from:'01-01',to:'12-31'}]"),
    ("location_lat",           "decimal",    "GPS latitude for solar calculations"),
    ("location_lng",           "decimal",    "GPS longitude for NASA POWER API"),
    ("location_name",          "varchar",    "Human-readable location label"),
]
for col, typ, desc in proj_cols:
    row = t.add_row()
    row.cells[0].text = col; row.cells[1].text = typ; row.cells[2].text = desc

doc.add_paragraph()
h3("buildings")
body("Foreign key to projects. Each building adds: area (m² for solar roof calculation), floors count, work_days, work_time_intervals, working_season_intervals, existing_solar_power.")

h3("floors")
body("Foreign key to buildings. Inherits or overrides work schedule from parent building.")

h3("rooms")
body("Foreign key to floors. Has name and optionally overrides work schedule.")

doc.add_paragraph()
h2("3.2 Component Tables")
body(
    "There are four component tables — one per hierarchy level. They share an identical "
    "column structure (only the foreign key differs). This design avoids a single generic "
    "'components' table and keeps foreign key constraints clean."
)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Column"; t.rows[0].cells[1].text = "Type"; t.rows[0].cells[2].text = "Meaning"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
comp_cols = [
    ("{level}_id",              "FK",         "Foreign key to parent (room_id / floor_id / building_id / project_id)"),
    ("component_type_id",       "FK",         "Link to component_types library (for name and defaults)"),
    ("power",                   "decimal",    "Rated apparent power in VA (volt-amperes) — NOT watts"),
    ("phases",                  "enum",       "'1phase' or '3phase'"),
    ("phase",                   "char(1)",    "Assigned phase: 'A', 'B', or 'C' (null = unassigned)"),
    ("power_factor",            "decimal",    "Component power factor 0.01–1.00 (default 1.0)"),
    ("quantity",                "int",        "Number of identical units"),
    ("group_name",              "varchar",    "Optional: loads in same group → only largest counts (N+1 diversity)"),
    ("priority",                "varchar",    "'critical', 'essential', or 'normal'"),
    ("needs_socket",            "boolean",    "Whether this component also needs a socket outlet counted"),
    ("usage_season",            "varchar",    "'all', 'summer', 'winter', 'spring', 'autumn'"),
    ("usage_day_type",          "varchar",    "'all', 'weekday', or 'weekend'"),
    ("usage_time_intervals",    "JSON array", "Daily on/off windows: [{start:'08:00',end:'18:00'}]"),
    ("is_motor",                "boolean",    "True → eligible for NEC/IEC 125% inrush sizing"),
]
for col, typ, desc in comp_cols:
    row = t.add_row()
    row.cells[0].text = col; row.cells[1].text = typ; row.cells[2].text = desc

doc.add_paragraph()
h2("3.3 Polymorphic Tables")
body(
    "Three tables use Laravel polymorphic relationships so the same model can belong to a project, "
    "building, floor, or room without four separate tables:"
)

h3("utility_lines")
bullet("lineable_type: 'App\\Models\\Project' | 'App\\Models\\Building' | etc.")
bullet("lineable_id: the parent entity's ID")
bullet("phases: '1phase' | '3phase'")
bullet("power: rated capacity in VA")
bullet("quantity: number of supply lines")

h3("generator_lines")
body("Identical structure to utility_lines but uses generable_type / generable_id. Stores diesel generator capacity.")

h3("sockets")
body("socketable_type / socketable_id. Stores phase_type (1phase/3phase), power per outlet (VA), quantity.")

doc.add_paragraph()
h2("3.4 Access Control Tables")
h3("project_users")
bullet("project_id (FK), user_id (FK), role: 'main' | 'normal'")
bullet("The project owner (projects.user_id) implicitly has 'admin' role and is NOT in this table.")
bullet("Roles: admin = full control, main = can edit, normal = read-only.")

h3("component_types")
body(
    "A library of reusable component templates. When a user adds a component, they pick a type; "
    "the defaults (power, power_factor, phases, usage_season, etc.) are pre-filled. "
    "is_preset=true entries are system-provided. is_motor=true flags motor loads for inrush calculation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — BACKEND
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Backend — Laravel Application")

# ── 4.1 AUTH ─────────────────────────────────────────────────────────────────
h2("4.1 Authentication & Authorization")

h3("Google OAuth 2.0 (Primary)")
body(
    "End users authenticate via Google. The frontend redirects to /auth/google/redirect. "
    "After Google consent, the callback at /auth/google/callback receives the user profile, "
    "creates or updates a User record, and issues a Sanctum personal access token. "
    "The token is returned in the JSON response and stored in the browser (localStorage or memory). "
    "Subsequent API requests send it as 'Authorization: Bearer {token}'."
)

h3("Admin Login (Secondary)")
body(
    "A separate credential-based login at POST /api/admin/login for system administrators. "
    "Admin users can view all users, activate/deactivate accounts, and delete users. "
    "Admin routes are protected by the 'admin' middleware which checks a flag on the User model."
)

h3("Role-Based Access Inside Projects")
body(
    "Every controller that touches project data calls $project->userRole($userId) before proceeding. "
    "This method returns 'admin' if the user owns the project, the stored role ('main'/'normal') "
    "if they are a member, or null if they have no access. Returning null or 'normal' on a "
    "write operation causes an immediate 403 response."
)
code_block(
    "public function userRole(int \$userId): ?string\n"
    "{\n"
    "    if ((int) \$this->user_id === \$userId) return 'admin';   // owner\n"
    "    \$member = \$this->projectUsers()->where('user_id', \$userId)->first();\n"
    "    return \$member?->role;  // 'main', 'normal', or null\n"
    "}"
)

# ── 4.2 PROJECT HIERARCHY ────────────────────────────────────────────────────
h2("4.2 Project Hierarchy")

body(
    "The system models a real building complex as a four-level tree. Each level is represented "
    "by a dedicated Eloquent model and database table."
)
code_block(
    "Project\n"
    "  └─ Building  (hasMany)\n"
    "       └─ Floor    (hasMany)\n"
    "            └─ Room     (hasMany)\n"
    "\n"
    "Each level also hasMany:\n"
    "  - {Level}Components  (dedicated component table)\n"
    "  - UtilityLine        (polymorphic via lineable)\n"
    "  - GeneratorLine      (polymorphic via generable)\n"
    "  - Socket             (polymorphic via socketable)"
)

body(
    "Work schedule attributes (work_days, work_time_intervals, working_season_intervals) cascade downward "
    "— if a floor does not define its own schedule, it inherits from its building, which inherits from "
    "the project. This is implemented explicitly in the controllers when building the component list."
)

h3("Schedule Inheritance Logic")
code_block(
    "\$pDays    = \$project->work_days;\n"
    "\$pSeasons = \$project->working_season_intervals;\n"
    "\n"
    "foreach (\$buildings as \$building) {\n"
    "    \$bDays = \$building->work_days ?? \$pDays;  // null → inherit\n"
    "    foreach (\$building->floors as \$floor) {\n"
    "        \$fDays = \$floor->work_days ?? \$bDays;\n"
    "        foreach (\$floor->rooms as \$room) {\n"
    "            \$rDays = \$room->work_days ?? \$fDays;\n"
    "        }\n"
    "    }\n"
    "}"
)

# ── 4.3 COMPONENTS ───────────────────────────────────────────────────────────
h2("4.3 Electrical Components System")

body(
    "A component represents one type of electrical load installed at a specific location in "
    "the hierarchy. Every component has a quantity — so 20 identical ceiling lights can be "
    "represented as a single record with quantity=20."
)

h3("The power Column Stores VA, NOT Watts")
body(
    "This is a critical design decision. The power column stores apparent power in volt-amperes (VA). "
    "This is the correct quantity for cable and breaker sizing. Active power (W) is derived by "
    "multiplying by the power factor (PF): P = S × PF. Reactive power Q is derived as: "
    "Q = P × tan(arccos(PF))."
)
code_block(
    "// Inside every accumulation loop:\n"
    "\$va = (float) \$component->power * (int) \$component->quantity;  // S = VA_rated × qty\n"
    "\$w  = \$va * \$pf;                                                // P = S × PF\n"
    "\$q  = \$pf < 1.0 ? \$w * tan(acos(min(1.0, \$pf))) : 0.0;       // Q = P·tan(arccos(PF))"
)

h3("Group Name — N+1 Redundancy Diversity")
body(
    "If multiple components in the same entity share the same group_name, only the one with "
    "the highest VA is counted in the optimized total. This models N+1 redundancy (e.g., two "
    "UPS units where only the larger runs while the other is hot-standby). The max_va path "
    "always counts all components regardless of group."
)

h3("Priority Classification")
body("Every component is tagged with one of three priority levels:")
t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text="Priority"; t.rows[0].cells[1].text="Examples"; t.rows[0].cells[2].text="Treatment"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
prio_rows = [
    ("critical","Servers, fire alarms, emergency lighting, network hubs",
     "DF=1.0 always. Runs 24h/day, 7 days/week. Never schedule-filtered."),
    ("essential","Core HVAC, medical equipment, key production machinery",
     "DF applied per hierarchy level. Schedule filtering applied."),
    ("normal","Workstation lighting, general HVAC, standard equipment",
     "DF applied per hierarchy level. Schedule filtering applied."),
]
for i, (p, ex, tr) in enumerate(prio_rows):
    t.rows[i+1].cells[0].text = p
    t.rows[i+1].cells[1].text = ex
    t.rows[i+1].cells[2].text = tr

# ── 4.4 POWER SOURCES ────────────────────────────────────────────────────────
h2("4.4 Power Sources — Utility, Generator, Sockets")

h3("Utility Lines")
body(
    "Represent the mains grid connection. Stored with power (VA) and phases. Multiple lines "
    "can exist per entity (e.g., a building with a primary and backup utility feed). "
    "In dispatch calculations the total utility capacity is sum(utility_lines.power) × 0.8 "
    "(applying a standard 0.8 power factor to convert VA to W for dispatch comparison)."
)

h3("Generator Lines")
body(
    "Same structure as utility lines. Represent diesel/gas generator sets. In dispatch they "
    "are the last resort after solar and utility are exhausted."
)

h3("Sockets — Demand Factor Model")
body(
    "Sockets (outlet points) are modelled separately from fixed components because their "
    "actual load is statistically unknown. The SocketDemandService applies IEC-derived "
    "demand factors based on outlet count:"
)
code_block(
    "First 10 outlets:       100% × 200 VA/outlet\n"
    "Next 10 outlets (11-20): 75% × 200 VA/outlet\n"
    "Beyond 20:              40% × 200 VA/outlet"
)
body(
    "At building and project level an additional coincidence factor is applied: "
    "CF = 1.00 for demand < 50 kVA, CF = 0.92 for 50–250 kVA, CF = 0.85 for > 250 kVA. "
    "This reflects the statistical impossibility of all outlets in a large building "
    "being loaded simultaneously."
)

# ── 4.5 TOTAL POWER CONTROLLER ────────────────────────────────────────────────
h2("4.5 TotalPowerController — The Core Power Calculation Engine")

body(
    "TotalPowerController is the most complex controller. It is called separately for each "
    "level (project, building, floor, room) and returns a comprehensive power report. "
    "All four methods share the same private helper functions."
)

h3("The sumPowerWithGroups Method")
body(
    "This is the single accumulation loop that populates BOTH the overall total vectors "
    "AND the per-priority (critical/essential/normal) buckets simultaneously. "
    "This ensures total_va and priority subtotals are always derived from the same source — "
    "eliminating any chance of inconsistency between them."
)
code_block(
    "private function sumPowerWithGroups(\n"
    "    \$query,\n"
    "    string \$entityKey,\n"
    "    bool  \$withPriority  = false,\n"
    "    float \$dfMultiplier  = 1.0     // IEC DF for this hierarchy level\n"
    "): array"
)
body("For each component in the query result:")
bullet("VA, W, Q are computed from power × PF.")
bullet("Max tracking (maxVA, maxW, maxQ) accumulates raw values with NO diversity factor — this feeds max_va.")
bullet("For critical priority: effectiveDf = 1.0 regardless of the passed dfMultiplier (critical loads are never diversified).")
bullet("For all other priorities: effectiveDf = dfMultiplier (the passed IEC cascade value).")
bullet("Diversified W and Q (wd, qd) are accumulated into wTotal, qTotal.")
bullet("The same wd/qd are accumulated into the correct priority bucket (critical/essential/normal).")
bullet("Group-max selection: within a group, only the component with the highest raw VA is kept.")

h3("IEC Diversity Factor Cascade")
body(
    "When the project() endpoint calls sumPowerWithGroups for each hierarchy level, "
    "it passes the compounded DF for that level's position in the hierarchy:"
)
code_block(
    "// project() endpoint:\n"
    "\$own      = sumPowerWithGroups(project->components,  dfMultiplier: 1.0)            // DF=1.0\n"
    "\$building = sumPowerWithGroups(BuildingComponent::,  dfMultiplier: DF_PROJECT)     // ×0.7\n"
    "\$floor    = sumPowerWithGroups(FloorComponent::,     dfMultiplier: DF_BUILDING×DF_PROJECT)  // ×0.56\n"
    "\$room     = sumPowerWithGroups(RoomComponent::,      dfMultiplier: DF_FLOOR×DF_BUILDING×DF_PROJECT) // ×0.504\n"
    "\n"
    "// Merge all into one combined $comp array:\n"
    "\$comp = mergePower(\$own, \$building, \$floor, \$room);"
)
note(
    "DF_FLOOR=0.9, DF_BUILDING=0.8, DF_PROJECT=0.7 are IEC 60364-8-1 constants. "
    "A room component's diversified contribution to the project total is × 0.504 of its nameplate VA."
)

h3("total_va Computation")
body(
    "After merging, total_va is computed from the combined P and Q vectors. Sockets are added "
    "to the total using the 0.95 assumed power factor (TARGET_POWER_FACTOR) but are NOT added "
    "to any priority bucket."
)
code_block(
    "\$sinPf   = sin(acos(0.95));                              // Q component for sockets\n"
    "\$totalW  = \$comp['w'] + \$socketDemand * 0.95;          // total active power (W)\n"
    "\$totalQ  = \$comp['q'] + \$socketDemand * \$sinPf;        // total reactive power (VAR)\n"
    "\$totalVa = round(sqrt(\$totalW**2 + \$totalQ**2), 2);    // total apparent power (VA)"
)
note(
    "VA is NEVER summed directly across components. Only P and Q are summed (they are additive). "
    "VA is always derived as sqrt(P²+Q²) at the very end. This is the only mathematically correct approach."
)

h3("max_va vs total_va vs priority_va")
t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text="Field"; t.rows[0].cells[1].text="What it means"; t.rows[0].cells[2].text="DF Applied"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
va_rows = [
    ("max_va",     "Sum of ALL components at nameplate VA × qty. Worst-case absolute maximum. Used for breaker sizing.", "None"),
    ("total_va",   "Diversified demand. Room components at ×0.504, floor at ×0.56, building at ×0.7. Realistic simultaneous demand.", "Yes (IEC cascade)"),
    ("critical_va","VA of critical components only, NO diversity factor (they run 24/7 at full load).", "None (always 1.0)"),
    ("essential_va / normal_va", "Priority subtotals with the same IEC DF as total_va.", "Yes (IEC cascade)"),
]
while len(t.rows) < len(va_rows) + 1:
    t.add_row()
for i, (f, m, d) in enumerate(va_rows):
    t.rows[i+1].cells[0].text=f; t.rows[i+1].cells[1].text=m; t.rows[i+1].cells[2].text=d

h3("Motor Inrush (NEC/IEC 125% Rule)")
body(
    "The single largest motor load (by per-unit VA) is identified across all hierarchy levels. "
    "Its contribution to max_va is increased by 25% (INRUSH_MULTIPLIER = 1.25) to account for "
    "locked-rotor inrush current at startup. This is added only to maxW and maxQ (not to the "
    "diversified totalW/totalQ), keeping it as a worst-case scenario indicator."
)
code_block(
    "\$delta = 0.25;  // extra 25% inrush\n"
    "\$addVa = motorPerUnitVA × motorQty × \$delta;\n"
    "\$addW  = \$addVa × \$motorPf;\n"
    "\$addQ  = \$addVa × sqrt(1 - \$motorPf²);\n"
    "\$maxW += \$addW;\n"
    "\$maxQ += \$addQ;"
)

# ── 4.6 LOAD PROFILE ────────────────────────────────────────────────────────
h2("4.6 LoadProfileController — Hourly Demand Curves")

body(
    "LoadProfileController builds a 24-element array (one value per hour 0–23) representing "
    "the total active power demand (kW) and reactive demand (kVAR) across the day. "
    "It applies the same IEC diversity factor cascade as TotalPowerController."
)

h3("Component Collection with DF")
body(
    "The addComponents() helper is called once per hierarchy level with the appropriate "
    "dfMultiplier. For each component, peak_w = VA × PF × qty × DF. The fmt() function "
    "stores this diversified peak_w along with the component's schedule metadata."
)
code_block(
    "// Room components get the full cascade:\n"
    "\$this->addComponents(\$result, \$room->components, 'room_id',\n"
    "    \$rDays, \$rSeasons,\n"
    "    self::DF_FLOOR * self::DF_BUILDING * self::DF_PROJECT  // = 0.504\n"
    ");"
)
body(
    "Exception: critical priority components bypass the dfMultiplier entirely (effectiveDf = 1.0) "
    "because critical loads are never diversified."
)

h3("computeHourlyW — Schedule Filtering")
body(
    "For each component in the prepared list, the computeHourlyW loop adds peak_w to every "
    "hour that falls within the component's usage_time_intervals. Components tagged as critical "
    "priority bypass the interval check entirely — their peak_w is added to all 24 hours, "
    "creating the flat baseline on the chart."
)
code_block(
    "if (\$c['priority'] === 'critical') {\n"
    "    for (\$h = 0; \$h < 24; \$h++) { \$hourlyP[\$h] += \$peakW; }\n"
    "    continue;  // no interval filtering needed\n"
    "}\n"
    "// Non-critical: only active during their intervals\n"
    "foreach (\$c['usage_time_intervals'] as \$interval) {\n"
    "    // ... add to hours between start and end\n"
    "}"
)
body("The same critical bypass applies in computeHourlyKvar for reactive power.")

h3("Solar Integration")
body(
    "After computing the hourly load arrays, the controller calls SolarIrradianceService "
    "to get a 24-element array of solar generation in kW. This is returned alongside the "
    "load data so the frontend can render both curves on the same chart."
)

# ── 4.7 SCHEDULE CONTROLLER ──────────────────────────────────────────────────
h2("4.7 ScheduleController — 24-hour Dispatch")

body(
    "ScheduleController generates the full Load Schedule page data. Unlike LoadProfileController "
    "(which does not filter by day type or season), ScheduleController builds two load profiles "
    "(max and optimized) for a specific month, day type (workday/weekend), and calendar day."
)

h3("Query Parameters")
t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text="Param"; t.rows[0].cells[1].text="Default"; t.rows[0].cells[2].text="Description"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
params = [
    ("month",    "current",  "1–12. Selects the solar irradiance for that month."),
    ("day_type", "workday",  "'workday', 'weekend', or 'all'. Filters component schedules."),
    ("day",      "15",       "Calendar day 1–31. Used for the NASA POWER API call to get actual daily GHI."),
]
for i, (p, d, desc) in enumerate(params):
    t.rows[i+1].cells[0].text=p; t.rows[i+1].cells[1].text=d; t.rows[i+1].cells[2].text=desc

doc.add_paragraph()
h3("Load Profile Construction")
body(
    "collectComponents() traverses the full hierarchy and produces a flat list of all "
    "components with their schedule metadata. Then buildHourlyW() applies day-type and "
    "season filters for each component and accumulates load into a 24-hour array. "
    "The 'optimized' mode also applies group-max: within each group only the highest-VA "
    "component is kept, simulating that only the primary unit runs (not the backup)."
)

h3("Solar Profile")
body(
    "The controller asks SolarIrradianceService for the 24-hour solar generation array for "
    "the specific calendar day. It first tries NASA POWER (real satellite data for that exact "
    "date in 2023). If the API is unavailable, it falls back to a static PSH lookup table "
    "with a sinusoidal bell-curve distribution."
)

h3("Source Dispatch")
body(
    "The SourceDispatchService.dispatch() method is called twice — once for the max load "
    "profile and once for the optimized profile. It applies a Solar → Utility → Generator "
    "priority order hour-by-hour, filling demand first with free solar, then the grid, "
    "then the generator. Any remaining unmet demand is recorded."
)

# ── 4.8 PHASE BALANCE ────────────────────────────────────────────────────────
h2("4.8 PhaseBalanceController — 3-Phase Optimization")

body(
    "Three-phase buildings receive their power across phases A, B, and C. Single-phase loads "
    "are connected to one phase. If the loads are unevenly distributed, current imbalance "
    "causes wasted neutral current, transformer heating, and voltage unbalance."
)

h3("Greedy Balancing Algorithm")
body(
    "The controller implements a classic greedy bin-packing approach to find the optimal "
    "phase assignment:"
)
bullet("1. Collect all single-phase load blocks (room components, floor components, socket demand). Each block is treated as a single atomic unit.")
bullet("2. Sort blocks in descending order by VA (largest first).")
bullet("3. For each block, assign it to whichever phase currently has the lowest total VA.")
bullet("4. This greedy-largest-first approach produces near-optimal balance for typical load distributions.")
code_block(
    "\$optVa = ['A' => 0.0, 'B' => 0.0, 'C' => 0.0];\n"
    "foreach (\$blocks as \$blk) {\n"
    "    \$ph = array_keys(\$optVa, min(\$optVa))[0];  // lowest phase\n"
    "    \$optVa[\$ph] += \$blk['va'];\n"
    "    // record assignment for this block\n"
    "}"
)

h3("Imbalance Metric")
body(
    "Phase imbalance is quantified as the percentage deviation between the maximum and minimum "
    "phase currents relative to the average. The threshold for warnings is >10%, and >20% is critical."
)
code_block(
    "// From imbalanceStatus():\n"
    "\$avg = (\$iA + \$iB + \$iC) / 3;\n"
    "\$imb = (max(\$iA,\$iB,\$iC) - min(\$iA,\$iB,\$iC)) / \$avg × 100;\n"
    "\n"
    "// Neutral current (Millman's formula approximation):\n"
    "\$iN = sqrt(max(0, \$iA² + \$iB² + \$iC² - \$iA×\$iB - \$iB×\$iC - \$iC×\$iA));"
)

h3("Apply Optimal Assignment")
body(
    "The applyOptimalBuilding endpoint writes the greedy-optimal phase assignment back to the "
    "database. It updates the phase column of all 1-phase components in each room, floor, "
    "and building according to the block assignment computed above."
)

# ── 4.9 SOLAR SERVICE ────────────────────────────────────────────────────────
h2("4.9 SolarIrradianceService — NASA POWER Integration")

h3("Primary Path: NASA POWER API")
body(
    "NASA's POWER (Prediction Of Worldwide Energy Resources) project provides satellite-derived "
    "solar irradiance data for any point on Earth. The service fetches hourly Global Horizontal "
    "Irradiance (GHI) in W/m² for a specific latitude, longitude, and date."
)
code_block(
    "GET https://power.larc.nasa.gov/api/temporal/hourly/point\n"
    "  ?parameters=ALLSKY_SFC_SW_DWN\n"
    "  &community=RE\n"
    "  &latitude=31.5129\n"
    "  &longitude=34.4581\n"
    "  &start=20230527\n"
    "  &end=20230527\n"
    "  &format=JSON\n"
    "  &time-standard=LST"
)
body(
    "Response contains keys like '2023052700', '2023052701', ... '2023052723' — the last two "
    "digits are the hour. Values are GHI in W/m². The null indicator value -999 is replaced with 0."
)
body(
    "Results are cached for 30 days using Laravel Cache (key: nasa_ghi_{lat}_{lng}_{month}_{day}). "
    "Historical irradiance data never changes, so 30-day caching is safe. The year 2023 is used "
    "as the historical reference year (current/future years return -999 null values)."
)

h3("Conversion to Panel Output Watts")
body(
    "Raw GHI (W/m²) is converted to panel power output using:"
)
code_block(
    "output_W = (GHI / 1000) × panelCapacity_W × performanceRatio\n"
    "\n"
    "// GHI / 1000 normalises to the STC irradiance (1000 W/m²)\n"
    "// panelCapacity_W = installed nameplate capacity\n"
    "// performanceRatio = 0.80 (accounts for inverter losses, wiring, temperature, soiling)"
)

h3("Fallback: Static PSH Lookup Table")
body(
    "When NASA is unavailable, the service falls back to a pre-computed table of Peak Sun Hours "
    "(PSH) indexed by absolute latitude (0°, 10°, 20°, 30°, 40°, 50°, 60°) and month. "
    "Southern hemisphere locations have their month index shifted by 6 (season flip). "
    "Intermediate latitudes are linearly interpolated."
)
body(
    "Solar declination is computed using Spencer's equation (accurate to ±0.5°). Sunrise/sunset "
    "times are derived from the hour-angle formula (assuming solar noon at 12:00). Generation is "
    "distributed as a sinusoidal bell curve between sunrise and sunset whose area equals "
    "the panel capacity × PSH × PR."
)

# ── 4.10 SOURCE DISPATCH ─────────────────────────────────────────────────────
h2("4.10 SourceDispatchService — Priority Dispatch")

body(
    "For each of the 24 hours, SourceDispatchService fills demand from three sources in strict "
    "priority order: Solar → Utility Grid → Generator. This mirrors the economic and operational "
    "reality: solar is free, grid is cheap, generator is expensive."
)
code_block(
    "for (\$h = 0; \$h < 24; \$h++) {\n"
    "    \$remaining = \$demand[\$h];\n"
    "    \$sU = min(\$solarAvail[\$h], \$remaining); \$remaining -= \$sU;\n"
    "    \$uU = min(\$utilityCapW,     \$remaining); \$remaining -= \$uU;\n"
    "    \$gU = min(\$genCapW,         \$remaining); \$remaining -= \$gU;\n"
    "    \$unmet[\$h] = max(0, \$remaining);\n"
    "}"
)
body("The stats output includes: solar kWh, utility kWh, generator kWh, unmet kWh, self-consumption ratio.")

# ── 4.11 SOCKET DEMAND ──────────────────────────────────────────────────────
h2("4.11 SocketDemandService — Outlet Demand Factors")

body(
    "Fixed components (servers, AC units, etc.) have known VA ratings. Socket outlets do not — "
    "we don't know what will be plugged in. The service applies IEC-derived demand factors:"
)
bullet("Assumes 200 VA per outlet as the base connected load.")
bullet("First 10 outlets: 100% utilisation.")
bullet("Outlets 11–20: 75% utilisation (not all medium-use outlets loaded simultaneously).")
bullet("Outlets 21+: 40% utilisation (large panels mostly idle).")
bullet("At building/project level, a coincidence factor (0.92–1.00) further reduces the aggregate.")

# ── 4.12 BACKUP ──────────────────────────────────────────────────────────────
h2("4.12 Backup & Restore System")

body(
    "The system provides full JSON serialization/deserialization of any entity in the hierarchy. "
    "A backup captures the entity plus all its descendants (buildings include all floors, rooms, "
    "and components; a room backup includes just its components and sockets)."
)
body("Backups can be:")
bullet("Downloaded as a JSON file to the user's browser (client-side backup).")
bullet("Stored on the server in the server_backups table (server-side backup).")
bullet("Used to restore/duplicate: restore replaces an entity's data; duplicate creates a new copy.")
body(
    "The duplicate feature is especially useful: a user can duplicate a building (including all "
    "its floors, rooms, and components) to quickly create a second identical building with different names."
)

# ── 4.13 PROJECT MEMBERS ────────────────────────────────────────────────────
h2("4.13 Project Member Management")

body(
    "The project owner (admin) can invite other users to collaborate on a project. "
    "Members are added by email address. Roles are: 'main' (can edit all data) or 'normal' (read-only). "
    "The owner always retains full admin control and cannot be removed."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — FRONTEND
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Frontend — React Application")

h2("5.1 Routing & Authentication")

body(
    "The frontend uses React Router v6. Routes are code-split with React.lazy() so each page "
    "loads its JavaScript bundle only when visited, keeping the initial load fast. "
    "A ProtectedRoute wrapper redirects unauthenticated users to the login page. "
    "An AuthProvider context stores the current user and token, providing them to all components."
)
code_block(
    "/ → DashboardPage (project list)\n"
    "/:projectName → ProjectPage\n"
    "/:projectName/:buildingName → BuildingPage\n"
    "/:projectName/:buildingName/:floorName → FloorPage\n"
    "/:projectName/:buildingName/:floorName/:roomName → RoomDetailPage\n"
    "/load-schedule/:projectId → LoadSchedulePage\n"
    "/phase-balance/:projectId → PhaseBalancePage\n"
    "/login → LoginPage\n"
    "/admin → AdminDashboardPage"
)

h2("5.2 Dashboard Page")
body(
    "Shows all projects the user owns or is a member of. Cards display the project name, "
    "building type, last updated time, and a role badge. "
    "From the dashboard, users can: create a new project, restore from a JSON backup file, "
    "restore from a server-side backup, manage project members."
)

h2("5.3 Project Page")
body(
    "The top-level view of a project. Shows: list of buildings, the PowerBanner with total/max "
    "VA and priority breakdown, the PowerSourcesBanner (solar/utility/generator sources), "
    "the ReactivePowerPanel, and the EntityComponents section for project-level components. "
    "From here users navigate into buildings, or go to Load Schedule / Phase Balance analysis pages."
)

h2("5.4 Building / Floor / Room Pages")
body(
    "Each level has a similar layout: the entity's own components (EntityComponents), "
    "sockets (EntitySockets), sources (PowerSourcesBanner), reactive power (ReactivePowerPanel), "
    "and the PowerBanner showing that level's diversified total. "
    "The floor page shows a grid of rooms. Each room card shows its total VA and a colour-coded "
    "priority badge."
)

h2("5.5 Load Schedule Page")
body(
    "The most visually rich page. Contains:"
)
bullet("Controls bar: month selector, mini calendar (day picker), Workday/Weekend/All Days toggle.")
bullet("Info pills: Sunrise/Sunset times, PSH (peak sun hours), Solar cap (installed capacity), NASA/static data source badge.")
bullet("Main chart: stacked area chart showing Demand (black outline), Solar (yellow), Utility (blue), Generator (orange), Unmet (red) for 24 hours.")
bullet("Three sub-charts: one each for Solar, Utility Grid, Generator showing individual source output vs the demand curve (dashed).")
bullet("Daily Energy Breakdown table: hours of use, kWh, and % share for each source. Shows 'All Load Covered' or 'Unmet Demand X kWh'.")
body(
    "The Load Schedule tab shows the per-component schedule (which components are active at which hours). "
    "The Sources tab allows editing utility and generator line capacities inline. "
    "The Optimized / Max Load toggle switches between the two dispatch profiles."
)

h2("5.6 Phase Balance Page")
body(
    "Displays the current and optimal phase distribution for each building. "
    "Shows a colour-coded imbalance indicator (green = balanced, yellow = warning, red = critical). "
    "The greedy-optimal assignment recommendation is shown alongside the current actual assignment. "
    "Users can click 'Apply Optimal' to write the recommended assignment to the database."
)

h2("5.7 PowerBanner Component")
body(
    "The main power summary strip visible at the top of every hierarchy page. Shows:"
)
bullet("MAX LOAD (kVA): the undiversified connected load.")
bullet("OPTIMIZED (kVA): the diversified simultaneous demand.")
bullet("Critical / Essential / Normal cards: each shows MAX and OPTIMIZED kVA for that priority tier.")
bullet("Sockets card: shows total outlet capacity and diversified socket demand.")
bullet("Toggle: VA / W mode to switch between apparent and active power display.")
bullet("kVAR indicator: total reactive power of the diversified load.")

h2("5.8 PowerSourcesBanner Component")
body(
    "Below the PowerBanner, this component allows users to configure and view power sources. "
    "Shows each utility line and generator line with their VA rating, phases, and computed W capacity. "
    "Inline edit forms allow modifying source parameters without navigating away. "
    "The solar section shows the computed PV capacity (from building roof areas) or user-declared "
    "existing capacity, along with the current solar_source mode."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. API Reference — All Endpoints")

note("All endpoints require Authorization: Bearer {token} unless noted. Base URL: /api")

h2("Authentication")
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text="Method & Path"; t.rows[0].cells[1].text="Description"; t.rows[0].cells[2].text="Auth"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
auth_ep = [
    ("GET  /auth/google/redirect",   "Redirect to Google OAuth",    "Public"),
    ("GET  /auth/google/callback",   "Receive OAuth token; return Sanctum token", "Public"),
    ("GET  /api/user",               "Get current authenticated user", "Required"),
    ("POST /api/logout",             "Revoke current token",         "Required"),
    ("POST /api/admin/login",        "Admin credential login",       "Public"),
]
for m, d, a in auth_ep:
    row = t.add_row(); row.cells[0].text=m; row.cells[1].text=d; row.cells[2].text=a

doc.add_paragraph()
h2("Project CRUD")
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text="Endpoint"; t.rows[0].cells[1].text="Description"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
proj_ep = [
    ("GET    /api/projects",                     "List all projects for current user"),
    ("POST   /api/projects",                     "Create new project"),
    ("GET    /api/projects/{project}",            "Get project detail"),
    ("PUT    /api/projects/{project}",            "Update project settings"),
    ("DELETE /api/projects/{project}",            "Delete project and all children"),
    ("GET    /api/projects/{project}/all-floors", "List all floors across all buildings"),
    ("GET    /api/projects/{project}/all-rooms",  "List all rooms across all floors"),
]
for ep, d in proj_ep:
    row = t.add_row(); row.cells[0].text=ep; row.cells[1].text=d

doc.add_paragraph()
h2("Power Analysis Endpoints")
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text="Endpoint"; t.rows[0].cells[1].text="Description"
for c in t.rows[0].cells: c.paragraphs[0].runs[0].bold = True
power_ep = [
    ("GET /api/projects/{project}/total-power",   "Full power report: total_va, max_va, priority breakdown, motor inrush, reactive power, capacitor bank"),
    ("GET /api/buildings/{building}/total-power", "Same for a single building scope"),
    ("GET /api/floors/{floor}/total-power",       "Same for a single floor scope"),
    ("GET /api/rooms/{room}/total-power",         "Same for a single room scope"),
    ("GET /api/projects/{project}/load-profile",  "24-hour hourly kW and kVAR load profile + solar curve"),
    ("GET /api/projects/{project}/schedule",      "Full dispatch schedule: load_max, load_optimized, solar, dispatch_max, dispatch_optimized, sunrise/sunset, PSH"),
    ("GET /api/projects/{project}/phase-balance", "Phase balance report for all buildings"),
    ("GET /api/buildings/{building}/phase-balance","Phase balance report for one building"),
    ("GET /api/floors/{floor}/phase-balance",     "Phase balance report for one floor"),
    ("POST /api/rooms/{room}/assign-phase",       "Manually assign phase A/B/C to all 1-phase components in a room"),
    ("POST /api/buildings/{building}/apply-optimal-phase", "Apply greedy-optimal phase assignment to entire building"),
]
for ep, d in power_ep:
    row = t.add_row(); row.cells[0].text=ep; row.cells[1].text=d

doc.add_paragraph()
h2("Schedule Endpoint Response Structure")
code_block(
    "{\n"
    "  month, month_name, day, day_type,\n"
    "  location: { lat, lng, name },\n"
    "  sunrise_hour, sunset_hour, peak_sun_hours,\n"
    "  solar_capacity_w, solar_data_source,\n"
    "  utility_capacity_va, generator_capacity_va,\n"
    "  load_max: [24 floats W],\n"
    "  load_optimized: [24 floats W],\n"
    "  solar: [24 floats W],\n"
    "  dispatch_max: {\n"
    "    solar_used, utility_used, generator_used, unmet,\n"
    "    stats: { solar_kwh, utility_kwh, generator_kwh, ... }\n"
    "  },\n"
    "  dispatch_optimized: { ... same structure ... },\n"
    "  hourly_kvar: [24 floats kVAR]\n"
    "}"
)

h2("Total Power Endpoint Response Structure")
code_block(
    "{\n"
    "  total_va, total (W), max_va, max_w,\n"
    "  critical_va, critical_w, critical_max_va, critical_max_w,\n"
    "  essential_va, essential_w, essential_max_va, essential_max_w,\n"
    "  normal_va, normal_w, normal_max_va, normal_max_w,\n"
    "  socket_demand_va, socket_connected_va,\n"
    "  own, own_va, building, building_va, floor, floor_va, room, room_va,\n"
    "  inrush_applied, inrush_component: { name, per_unit_va, quantity, ... },\n"
    "  total_kvar, max_kvar, system_power_factor,\n"
    "  pf_correction_recommended,\n"
    "  capacitor_bank_kvar, capacitor_bank_uf, capacitor_bank_target_pf,\n"
    "  current_before_correction_a, current_after_correction_a,\n"
    "  current_reduction_percent, correction_note,\n"
    "  solar_computed\n"
    "}"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ELECTRICAL ENGINEERING THEORY
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Electrical Engineering Theory")

h2("7.1 Apparent / Active / Reactive Power & Power Factor")

body(
    "AC electrical systems carry three types of power simultaneously. Understanding the "
    "distinction is critical because cables, breakers, and transformers are sized on apparent "
    "power (VA), while energy bills are based on active power (kWh)."
)
t = doc.add_table(rows=4, cols=4)
t.style = 'Table Grid'
for i, h in enumerate(["Quantity", "Symbol", "Unit", "Description"]):
    t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
pw_rows = [
    ("Apparent Power", "S", "VA (volt-amperes)", "Total power the supply must deliver. S = V × I"),
    ("Active Power",   "P", "W (watts)",         "Real work done. P = S × PF = V × I × cos(φ)"),
    ("Reactive Power", "Q", "VAR (volt-ampere reactive)", "Energy stored and released by inductors/capacitors. Q = P × tan(arccos(PF))"),
]
for i, (qty, sym, unit, desc) in enumerate(pw_rows):
    t.rows[i+1].cells[0].text=qty; t.rows[i+1].cells[1].text=sym
    t.rows[i+1].cells[2].text=unit; t.rows[i+1].cells[3].text=desc

doc.add_paragraph()
body("Power Triangle: S² = P² + Q²  →  S = √(P² + Q²)")
body(
    "Power Factor (PF) = P / S = cos(φ). A PF of 1.0 means all power does useful work. "
    "Motors, transformers, and fluorescent ballasts cause PF < 1 (lagging). A low PF means "
    "the supply carries more current than necessary, wasting capacity."
)
note(
    "In this system, P and Q are ALWAYS summed as scalars (they are additive for loads sharing "
    "the same frequency). VA is NEVER summed directly — it is computed as sqrt(P²+Q²) at the end. "
    "Summing VAs would be mathematically incorrect because it ignores phase angle differences."
)

h2("7.2 IEC 60364-8-1 Diversity Factors")

body(
    "In a large building, not every electrical load operates at the same time. A university's "
    "classrooms are not all occupied at 100% capacity simultaneously. IEC 60364-8-1 defines "
    "Diversity Factors (DF) to account for this statistical reality:"
)
t = doc.add_table(rows=5, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Hierarchy Level", "DF Value", "Interpretation"]):
    t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
df_rows = [
    ("Room level",     "1.0",   "Leaf node — no diversity reduction within a room"),
    ("Floor level",    "0.9",   "10% reduction when aggregating rooms to floor panel"),
    ("Building level", "0.8",   "20% reduction when aggregating floors to building panel"),
    ("Project level",  "0.7",   "30% reduction when aggregating buildings to project intake"),
]
for i, (l, v, d) in enumerate(df_rows):
    t.rows[i+1].cells[0].text=l; t.rows[i+1].cells[1].text=v; t.rows[i+1].cells[2].text=d

doc.add_paragraph()
body("Cascaded effective DF for a room component viewed at project level:")
code_block("Effective DF = DF_FLOOR × DF_BUILDING × DF_PROJECT = 0.9 × 0.8 × 0.7 = 0.504")
body(
    "A room component rated at 1000 VA contributes only 504 VA to the project total. "
    "This is why total_va ≈ 50% of max_va in a typical multi-building complex."
)
note(
    "Critical priority components are EXEMPT from diversity factors by definition — they must "
    "be on at all times, so their full nameplate VA counts even at the project level."
)

h2("7.3 Motor Inrush Current — NEC/IEC 125% Rule")

body(
    "Electric motors draw 6–8× their rated current for the first few cycles after startup "
    "(locked-rotor condition). Once running, current drops to rated value. This inrush current "
    "does not last long enough to trip a thermal breaker, but the electrical supply and cables "
    "must be sized to carry it."
)
body(
    "NEC Article 430 and IEC 60947 require that the largest motor in a panel be sized at "
    "125% of its rated full-load current. The system identifies the single highest-VA motor "
    "across all hierarchy levels and adds 25% of its VA to the max_va calculation:"
)
code_block(
    "sized_VA = motor_rated_VA × 1.25\n"
    "inrush_addition = motor_rated_VA × 0.25  // the extra 25%\n"
    "\n"
    "// Only applied to max_va (worst case), not to total_va (operational)"
)

h2("7.4 Reactive Power Correction & Capacitor Bank Sizing")

body(
    "When system PF is below the target (PENRA standard: PF_target = 0.95), a capacitor bank "
    "can inject leading reactive current to cancel the lagging reactive current from inductive "
    "loads. This raises the effective PF, reducing current draw and freeing up transformer capacity."
)

h3("Capacitor Bank Sizing Formula")
code_block(
    "Q_correction = Q_actual - P × tan(arccos(PF_target))\n"
    "\n"
    "// Round up to nearest 0.5 kVAR step (standard bank increment):\n"
    "bank_kVAR = ceil(Q_correction / 500) × 0.5\n"
    "\n"
    "// For Delta (Δ) configuration — each capacitor sees full line voltage (400 V):\n"
    "C_phase = (bank_kVAR × 1000 / 3) / (2π × 50 × 400²)   [Farads]\n"
    "C_phase_μF = C_phase × 1,000,000"
)
body(
    "The system always recommends Delta (Δ) connection because it allows each capacitor to "
    "see the full 400 V line-to-line voltage, requiring smaller capacitance values. "
    "Star (Y) connection would require the same kVAR at only 230 V (1/√3 times smaller), "
    "meaning larger capacitance values."
)

h3("Line Current Before and After Correction")
code_block(
    "// 3-phase system:\n"
    "I_before = S_before / (√3 × 400)\n"
    "I_after  = S_after  / (√3 × 400)\n"
    "\n"
    "// Current reduction %:\n"
    "reduction = (I_before - I_after) / I_before × 100"
)

h2("7.5 Phase Imbalance & Neutral Current")

body(
    "In a balanced 3-phase system, the neutral conductor carries near-zero current because the "
    "three phase currents sum to zero (120° apart). When loads are unequal across phases, "
    "the neutral current grows. Excessive neutral current causes:"
)
bullet("Overheating of the neutral conductor (which is often undersized relative to phase conductors).")
bullet("Voltage unbalance — phases with more load have lower voltage.")
bullet("Increased transformer losses and potential winding damage.")

body("The neutral current approximation used in this system:")
code_block(
    "// For resistive loads (PF=1 assumed in this approximation):\n"
    "I_N = √(I_A² + I_B² + I_C² - I_A×I_B - I_B×I_C - I_C×I_A)\n"
    "\n"
    "// This is derived from the phasor sum |I_A∠0° + I_B∠120° + I_C∠240°|"
)
note(
    "IEC 60038 recommends phase imbalance < 3% in distribution systems, < 2% in "
    "sensitive loads. This system uses 10% as a warning threshold and 20% as critical, "
    "which are appropriate for building-level analysis (not transmission-level)."
)

h2("7.6 Solar PV — PSH, STC, Performance Ratio")

h3("Standard Test Conditions (STC)")
body(
    "Solar panel nameplate ratings (e.g., '100W panel') are measured at STC: "
    "irradiance = 1000 W/m², cell temperature = 25°C, AM1.5 spectrum. "
    "Real-world conditions are almost always worse (higher temperature, lower irradiance), "
    "so actual output is typically 10–25% below nameplate."
)

h3("Peak Sun Hours (PSH)")
body(
    "PSH is the equivalent number of hours per day at 1000 W/m² irradiance that would deliver "
    "the same total daily solar energy as the actual variable irradiance profile. "
    "Example: Gaza in May has PSH ≈ 7.3 h/day. A 100 kW system produces 100 × 7.3 = 730 kWh/day "
    "before losses."
)

h3("Performance Ratio (PR)")
body(
    "PR accounts for all real-world losses: inverter efficiency (~97%), wiring resistance (~1%), "
    "temperature de-rating (~5% for hot climates), soiling/dust (~3%), shading (~2%). "
    "Typical PR = 0.75–0.85. This system uses PR = 0.80."
)
code_block(
    "actual_daily_kWh = panel_capacity_kW × PSH × PR\n"
    "\n"
    "// Hourly NASA path:\n"
    "hourly_output_W = (GHI_W_m2 / 1000) × capacity_W × PR\n"
    "\n"
    "// Roof area estimation (when solar_source = 'max'):\n"
    "solar_capacity_W = building_area_m2 × 0.17 × 1000 × 0.75\n"
    "// 17% panel coverage of roof area, 1000 W/m² STC, 0.75 packing factor"
)

h2("7.7 Priority-Based Load Classification")

body(
    "Electrical loads in a facility are classified by their criticality to operations. This "
    "determines UPS sizing, generator sizing, and load-shedding order during outages:"
)
t = doc.add_table(rows=4, cols=4)
t.style = 'Table Grid'
for i, h in enumerate(["Priority", "Examples", "UPS Backed?", "Treatment in System"]):
    t.rows[0].cells[i].text=h; t.rows[0].cells[i].paragraphs[0].runs[0].bold=True
pri_rows = [
    ("Critical", "Servers, fire/security systems, ICU equipment, telecom hubs",
     "Yes — zero transfer time", "DF=1.0, always on 24h/7d, flat baseline on chart"),
    ("Essential", "Core HVAC, main lab equipment, key production lines",
     "Often — seconds tolerance", "IEC DF applied, schedule filtering applied"),
    ("Normal", "Office lighting, general HVAC, standard workstations",
     "No — shed first", "IEC DF applied, schedule filtering applied"),
]
for i, (p, ex, ups, tr) in enumerate(pri_rows):
    t.rows[i+1].cells[0].text=p; t.rows[i+1].cells[1].text=ex
    t.rows[i+1].cells[2].text=ups; t.rows[i+1].cells[3].text=tr

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — CALCULATION WALKTHROUGHS
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Calculation Walkthroughs — End-to-End Examples")

h2("8.1 Example: Computing total_va for a Project")

body("Scenario: Islamic University — 2 buildings, multiple floors and rooms.")
body("Step 1 — Collect all components at each level and apply DF:")
code_block(
    "Project-own components:     P=0W    (no project-level components)\n"
    "Building components × 0.7:  P = 0W  (no building-level components)\n"
    "Floor components × 0.56:    P_floor = (floor crit 1.7kVA×PF + floor ess 2.42kVA×PF) × 0.56\n"
    "Room components × 0.504:    P_room  = (room ess 24.56kVA×PF + room norm 4.44kVA×PF) × 0.504\n"
    "\n"
    "Critical components (floor): DF overridden to 1.0\n"
    "  critical_P = 1.7kVA × PF × 1.0 (not ×0.56)"
)
body("Step 2 — Add socket demand:")
code_block(
    "socket_demand = 30.79 kVA\n"
    "socket_P = 30.79 × 0.95 = 29.25 kW\n"
    "socket_Q = 30.79 × sin(acos(0.95)) = 9.63 kVAR"
)
body("Step 3 — Compute final VA:")
code_block(
    "total_P = component_P_diversified + socket_P\n"
    "total_Q = component_Q_diversified + socket_Q\n"
    "total_VA = sqrt(total_P² + total_Q²) ≈ 50.27 kVA"
)
body("Step 4 — Verify consistency:")
code_block(
    "critical_VA  = 1.70 kVA   (no DF, 24h/7d)\n"
    "essential_VA = 13.73 kVA  (DF applied)\n"
    "normal_VA    = 5.34 kVA   (DF applied)\n"
    "sum of priority VAs       = 20.77 kVA\n"
    "plus socket demand        = 30.79 kVA\n"
    "subtotal sum              ≈ 51.56 kVA\n"
    "total_VA                  = 50.27 kVA\n"
    "gap (vector geometry)     = 1.29 kVA = 2.5%  ✓ (< 5% = acceptable)"
)

h2("8.2 Example: Capacitor Bank Sizing")
code_block(
    "total_W  = 49.3 kW   (active power)\n"
    "total_Q  = 9.61 kVAR (reactive power, from system)\n"
    "total_VA = 50.27 kVA\n"
    "system_PF = 49.3 / 50.27 = 0.981  → above 0.85 threshold → no correction needed\n"
    "\n"
    "Hypothetical example if PF were 0.78:\n"
    "  Q_actual    = P × tan(arccos(0.78)) = 49.3 × 0.802 = 39.5 kVAR\n"
    "  Q_target    = P × tan(arccos(0.95)) = 49.3 × 0.329 = 16.2 kVAR\n"
    "  Q_cap       = 39.5 - 16.2 = 23.3 kVAR\n"
    "  bank_kVAR   = ceil(23.3 / 0.5) × 0.5 = 23.5 kVAR\n"
    "  C_phase     = (23500/3) / (2π × 50 × 400²) = 247 μF/phase (Delta)"
)

h2("8.3 Example: Solar Dispatch for a Day")
code_block(
    "Location: Gaza (31.51°N), May 27 (workday), capacity: 100 kW\n"
    "\n"
    "NASA POWER returns GHI data for 2023-05-27:\n"
    "  Hour 06: 42 W/m²  → output = (42/1000) × 100,000 × 0.80 = 3,360 W\n"
    "  Hour 10: 840 W/m² → output = 67,200 W\n"
    "  Hour 13: 950 W/m² → output = 76,000 W\n"
    "  Hour 18: 12 W/m²  → output = 960 W\n"
    "\n"
    "Dispatch at hour 10 (load_optimized = 28,000 W):\n"
    "  solar_used    = min(67200, 28000) = 28,000 W  → demand fully covered by solar\n"
    "  utility_used  = 0 W\n"
    "  generator_used = 0 W\n"
    "\n"
    "Dispatch at hour 17 (load = 28,000 W, solar dropping to 25,000 W):\n"
    "  solar_used    = 25,000 W\n"
    "  remaining     = 3,000 W\n"
    "  utility_used  = 3,000 W  (utility cap >> 3000 W)\n"
    "  generator_used = 0 W"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — KNOWN LIMITATIONS & FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Known Limitations & Future Work")

h2("9.1 Current Limitations")
bullet("Solar profile uses a single day per month from NASA (year 2023). Weather variability (clouds, storms) within a month is not modelled.")
bullet("ScheduleController does not apply IEC diversity factors — it computes raw hour-by-hour load sums. This means the schedule chart load peak may differ from TotalPowerController's diversified total.")
bullet("Phase balance neutral current formula assumes resistive loads (PF=1). Inductive loads introduce phase shift that modifies the true neutral current calculation.")
bullet("Generator fuel consumption is not modelled — the system only shows hours and kWh, not litres of diesel or cost.")
bullet("No battery storage model — current dispatch goes directly from solar to grid to generator with no storage buffer.")
bullet("The system is single-location (one GPS coordinate per project). Multi-site projects with different weather profiles are not supported.")
bullet("Component library (component_types) is predefined; users cannot add custom types.")

h2("9.2 Future Work")
bullet("Battery storage model: add kWh battery capacity, charge/discharge logic in dispatch, and battery SoC tracking across 24 hours.")
bullet("Multi-day and multi-season dispatch: run the dispatch model for an entire year and report annual energy costs, peak months, generator runtime.")
bullet("Cost analysis: add electricity tariff input, fuel price, and compute annual cost comparison between solar+grid vs generator+grid scenarios.")
bullet("Export to IEC standard single-line diagram (SLD) PDF.")
bullet("Real-time monitoring integration: if the facility has smart meters, stream actual consumption against the predicted profile.")
bullet("Mobile-responsive layout for on-site use by electricians.")
bullet("Harmonics analysis for non-linear loads (VFDs, switched-mode power supplies).")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = r"e:\graduation project\power-profile\PROJECT_REPORT.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
